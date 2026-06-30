#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate.py — Génère les .docx des artefacts de cadrage EduTutor IA.

Convertit les fichiers Markdown du dossier cadrage (personas.md, et chaque
perturbations/*.md) en documents Word soignés (.docx placés à côté du .md) :
page de garde, styles de titres, listes et tableaux.

Dépendance : python-docx  (pip install python-docx)
Usage      : python generate.py

[Note pédagogique] La « source de vérité » reste le Markdown (versionnable,
diffable) ; le .docx est un rendu régénérable, pas l'original.
"""
import os
import re
import glob

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
SLATE = RGBColor(0x33, 0x41, 0x55)

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`.+?`)")
# Liens markdown [texte](url) -> on ne garde que le texte
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def strip_links(text):
    return LINK_RE.sub(r"\1", text)


def add_runs(paragraph, text):
    """Ajoute du texte avec gras (**), italique (*) et code (`)."""
    text = strip_links(text)
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1]); r.italic = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1]); r.font.name = "Consolas"
        else:
            paragraph.add_run(token)


def cover_page(doc, title):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EduTutor IA"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = INDIGO
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cadrage produit"); r.font.size = Pt(16); r.font.color.rgb = SLATE
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(22)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Livrable Jour 1 — APOCAL'IPSSI 2026"); r.italic = True; r.font.color.rgb = SLATE
    doc.add_page_break()


def flush_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header = cells[0]
    body = cells[2:] if len(cells) > 2 else []
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(strip_links(h).replace("**", "")); run.bold = True
    for row in body:
        tr = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(tr):
                tr[i].paragraphs[0].text = ""
                add_runs(tr[i].paragraphs[0], val)
    doc.add_paragraph()


def render(md_path, docx_path):
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = "Document"
    for l in lines:
        if l.startswith("# "):
            title = l[2:].strip(); break
    cover_page(doc, title)

    table_buf = []
    first_h1_seen = False
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("|"):
            table_buf.append(line)
            continue
        elif table_buf:
            flush_table(doc, table_buf); table_buf = []

        if line.startswith("# "):
            if not first_h1_seen:
                first_h1_seen = True
            else:
                doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\s*\d+\.\s+", "", line))
        elif line.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
            add_runs(p, line.lstrip("> ").strip())
        elif line.strip() in ("---", "***", "___"):
            pass
        elif line.strip() == "":
            pass
        else:
            add_runs(doc.add_paragraph(), line)

    if table_buf:
        flush_table(doc, table_buf)

    doc.save(docx_path)
    return docx_path


# Convention APOCAL'IPSSI : equipe-6-<slug>-vY.Z.docx. La source de vérité
# reste le .md ; les .docx sont des rendus régénérables et auto-découverts.
# Slug de sortie par fichier source (sinon le nom du fichier est réutilisé).
SLUGS = {
    "03_customer_journey_map": "customer-journey-map",
}
# Version par artefact (v1.0 par défaut).
VERSIONS = {
    "personas": "v1.2",
}


def out_name(stem):
    slug = SLUGS.get(stem, stem)
    version = VERSIONS.get(stem, "v1.0")
    return f"equipe-6-{slug}-{version}.docx"


def main():
    n = 0
    # Tous les artefacts de cadrage (auto-découverts) -> nommage equipe-6
    for md in sorted(glob.glob(os.path.join(HERE, "*.md"))):
        stem = os.path.splitext(os.path.basename(md))[0]
        out = os.path.join(HERE, out_name(stem))
        render(md, out)
        print(f"[OK] {os.path.relpath(out, HERE)}")
        n += 1
    # Perturbations (préfixe equipe-6 ajouté au nom du fichier)
    for md in sorted(glob.glob(os.path.join(HERE, "perturbations", "*.md"))):
        base = os.path.splitext(os.path.basename(md))[0]
        out = os.path.join(os.path.dirname(md), f"equipe-6-{base}.docx")
        render(md, out)
        print(f"[OK] {os.path.relpath(out, HERE)}")
        n += 1
    print(f"\n{n} document(s) genere(s).")


if __name__ == "__main__":
    main()
